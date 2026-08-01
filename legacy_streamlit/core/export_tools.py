from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

from docx import Document


def conversation_to_markdown(messages: list[dict[str, Any]]) -> str:
    lines = ["# Kimi Workspace conversation", ""]
    for message in messages:
        speaker = "You" if message.get("role") == "user" else "Kimi"
        lines.append(f"## {speaker}")
        for attachment in message.get("attachments", []):
            lines.append(f"- Attachment: {attachment.get('name', 'file')}")
        lines.append(message.get("text", ""))
        for index, source in enumerate(message.get("sources", []), start=1):
            lines.append(
                f"[{index}] {source.get('title', 'Source')}: {source.get('url', '')}"
            )
        lines.append("")
    return "\n".join(lines)


def conversation_to_json(messages: list[dict[str, Any]]) -> str:
    safe_messages: list[dict[str, Any]] = []
    for message in messages:
        safe = {
            "role": message.get("role"),
            "text": message.get("text", ""),
            "created_at": message.get("created_at"),
            "sources": message.get("sources", []),
            "attachments": [
                {
                    "name": item.get("name"),
                    "kind": item.get("kind"),
                    "summary": item.get("summary"),
                }
                for item in message.get("attachments", [])
            ],
        }
        safe_messages.append(safe)
    return json.dumps(
        {
            "exported_at": datetime.utcnow().isoformat() + "Z",
            "messages": safe_messages,
        },
        ensure_ascii=False,
        indent=2,
    )


def answer_to_docx(answer: str, title: str = "Kimi response") -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for block in answer.split("\n\n"):
        if block.strip():
            document.add_paragraph(block.strip())
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
