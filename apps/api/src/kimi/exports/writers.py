"""Artifact writers: DOCX, XLSX, Markdown, JSON, CSV.

The Word writer is the one the audit singled out. The prototype's
``answer_to_docx`` split on blank lines and wrote flat paragraphs, so headings
arrived as literal ``## Heading``, lists became run-on prose, tables became pipe
soup, and — although the system prompt mandates ``[1]`` citations — the exported
file contained dangling markers with no bibliography at all.
"""

from __future__ import annotations

import csv
import io
import json
import re
from datetime import UTC, datetime
from typing import Any

from kimi.exports.markdown import BlockKind, Span, is_rtl, parse_markdown

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

_SAFE_NAME = re.compile(r"[^\w\-. ]+", re.UNICODE)


def safe_stem(text: str, fallback: str = "kimi-export") -> str:
    """A stable, filesystem-safe stem derived from a title."""
    stem = _SAFE_NAME.sub("", (text or "").strip()).strip().replace(" ", "-")
    stem = re.sub(r"-{2,}", "-", stem).strip("-.")
    return (stem[:60] or fallback).lower()


def timestamped(stem: str, extension: str) -> str:
    return f"{stem}-{datetime.now(UTC):%Y%m%d-%H%M}.{extension}"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Insert a real clickable hyperlink.

    python-docx has no public API for this, so the relationship and the w:hyperlink
    element are built directly. Without it, exported citations are dead text.
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement  # type: ignore[attr-defined]

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")

    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0171DD")
    props.append(colour)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)

    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _set_rtl(paragraph: Any) -> None:
    """Mark a paragraph right-to-left so Arabic exports read correctly."""
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement  # type: ignore[attr-defined]

    props = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    props.append(bidi)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def _write_spans(paragraph: Any, spans: list[Span]) -> None:
    for span in spans:
        if span.href:
            _add_hyperlink(paragraph, span.href, span.text)
            continue
        run = paragraph.add_run(span.text)
        run.bold = span.bold
        run.italic = span.italic
        if span.code:
            run.font.name = "Menlo"


def answer_to_docx(
    *,
    title: str,
    body_markdown: str,
    sources: list[dict[str, Any]] | None = None,
    subtitle: str = "",
) -> bytes:
    """Render an answer as a Word document with real structure."""
    from docx import Document
    from docx.shared import Pt

    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)

    heading = document.add_heading(title or "Kimi Workspace", level=0)
    if is_rtl(title):
        _set_rtl(heading)

    if subtitle:
        meta = document.add_paragraph(subtitle)
        meta.runs[0].italic = True

    rtl_document = is_rtl(body_markdown)

    for block in parse_markdown(body_markdown):
        if block.kind is BlockKind.HEADING:
            paragraph = document.add_heading(level=min(block.level, 4))
            _write_spans(paragraph, block.spans)
        elif block.kind is BlockKind.BULLET:
            paragraph = document.add_paragraph(style="List Bullet")
            _write_spans(paragraph, block.spans)
        elif block.kind is BlockKind.NUMBERED:
            paragraph = document.add_paragraph(style="List Number")
            _write_spans(paragraph, block.spans)
        elif block.kind is BlockKind.QUOTE:
            paragraph = document.add_paragraph(style="Intense Quote")
            _write_spans(paragraph, block.spans)
        elif block.kind is BlockKind.CODE:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Menlo"
            run.font.size = Pt(9)
            continue
        elif block.kind is BlockKind.RULE:
            document.add_paragraph("―" * 24)
            continue
        elif block.kind is BlockKind.TABLE:
            if not block.rows:
                continue
            table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(block.rows):
                for c, cell in enumerate(row):
                    if c >= len(table.columns):
                        break
                    target = table.cell(r, c).paragraphs[0]
                    _write_spans(target, cell)
                    if r == 0:
                        for run in target.runs:
                            run.bold = True
            document.add_paragraph()
            continue
        else:
            paragraph = document.add_paragraph()
            _write_spans(paragraph, block.spans)

        if rtl_document:
            _set_rtl(paragraph)

    # The bibliography the prototype never wrote, despite mandating [n] markers.
    if sources:
        document.add_page_break()  # type: ignore[no-untyped-call]
        document.add_heading("Sources", level=1)
        for source in sources:
            index = source.get("index", "")
            paragraph = document.add_paragraph()
            paragraph.add_run(f"[{index}] ").bold = True
            paragraph.add_run(str(source.get("title") or source.get("url") or "Untitled"))

            details = document.add_paragraph()
            bits = [
                str(source.get("publisher") or ""),
                str(source.get("published_at") or "")[:10],
                str(source.get("status_label") or ""),
            ]
            run = details.add_run(" · ".join(b for b in bits if b))
            run.italic = True
            run.font.size = Pt(9)

            url = str(source.get("url") or "")
            if url:
                link_paragraph = document.add_paragraph()
                _add_hyperlink(link_paragraph, url, url)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def rows_to_xlsx(sheets: dict[str, list[list[Any]]]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    book = Workbook()
    book.remove(book.active)
    for name, rows in sheets.items():
        sheet = book.create_sheet(title=(name or "Sheet")[:31])
        for row in rows:
            sheet.append(row)
        if rows:
            for cell in sheet[1]:
                cell.font = Font(bold=True)
            for column in range(1, len(rows[0]) + 1):
                width = max(
                    (len(str(r[column - 1])) for r in rows[:200] if column - 1 < len(r)),
                    default=10,
                )
                sheet.column_dimensions[sheet.cell(1, column).column_letter].width = min(
                    max(width + 2, 10), 60
                )
    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()


def sources_to_xlsx(sources: list[dict[str, Any]]) -> bytes:
    header = [
        "#",
        "Title",
        "Publisher",
        "Published",
        "Date verified",
        "Status",
        "Provider",
        "Retrieval",
        "URL",
    ]
    rows: list[list[Any]] = [header]
    for source in sources:
        rows.append(
            [
                source.get("index", ""),
                source.get("title", ""),
                source.get("publisher", ""),
                (source.get("published_at") or "")[:19],
                "yes" if source.get("date_verified") else "no",
                source.get("status_label", ""),
                source.get("provider", ""),
                source.get("retrieval", ""),
                source.get("url", ""),
            ]
        )
    return rows_to_xlsx({"Sources": rows})


# ---------------------------------------------------------------------------
# Text formats
# ---------------------------------------------------------------------------


def conversation_to_markdown(
    *, title: str, messages: list[dict[str, Any]], generated_at: datetime | None = None
) -> str:
    when = (generated_at or datetime.now(UTC)).isoformat()
    lines = [f"# {title}", "", f"*Exported {when}*", ""]

    for message in messages:
        role = str(message.get("role", "")).lower()
        speaker = {"user": "You", "assistant": "Kimi"}.get(role, role.title())
        lines.append(f"## {speaker}")
        lines.append("")
        lines.append(str(message.get("content") or "").strip() or "*(no content)*")
        lines.append("")

        for citation in message.get("citations") or []:
            lines.append(
                f"> [{citation.get('index')}] {citation.get('title')} — "
                f"{citation.get('publisher')} — {citation.get('url')}"
            )
        if message.get("citations"):
            lines.append("")
    return "\n".join(lines)


def conversation_to_json(
    *, title: str, messages: list[dict[str, Any]], generated_at: datetime | None = None
) -> str:
    """Round-trippable export.

    Unlike the prototype's, this keeps every field it declares. That export
    dropped ``image_base64`` while keeping ``kind: "image"``, so re-importing it
    raised KeyError on every rerun and wedged the whole app (AUDIT §5).
    """
    payload = {
        "schema": "kimi.conversation/1",
        "title": title,
        # datetime.utcnow() is deprecated and naive; this is timezone-aware.
        "exported_at": (generated_at or datetime.now(UTC)).isoformat(),
        "messages": messages,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def sources_to_csv(sources: list[dict[str, Any]]) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(
        [
            "index",
            "title",
            "publisher",
            "published_at",
            "date_verified",
            "status",
            "provider",
            "retrieval",
            "url",
        ]
    )
    for source in sources:
        writer.writerow(
            [
                source.get("index", ""),
                source.get("title", ""),
                source.get("publisher", ""),
                source.get("published_at", ""),
                source.get("date_verified", ""),
                source.get("status_label", ""),
                source.get("provider", ""),
                source.get("retrieval", ""),
                source.get("url", ""),
            ]
        )
    return buffer.getvalue()
